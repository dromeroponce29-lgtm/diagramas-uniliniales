"""
Exportador DOCX profesional con python-docx.
Genera documento de especificaciones técnicas con portada, índice y
una sección por partida.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


COLOR_PRIMARY = RGBColor(0x0F, 0x2C, 0x4A)


def _aplicar_estilos(doc: Document) -> None:
    # Estilo Normal
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(18)
    h1.font.color.rgb = COLOR_PRIMARY
    h1.font.bold = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(14)
    h2.font.color.rgb = COLOR_PRIMARY
    h2.font.bold = True

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.color.rgb = COLOR_PRIMARY


def _agregar_contenido_md(doc: Document, contenido: str) -> None:
    """
    Convierte Markdown simple a párrafos de Word.
    Soporta headings ##, listas con - o *, negrita **texto**.
    """
    for linea in (contenido or "").split("\n"):
        l = linea.rstrip()
        if not l.strip():
            doc.add_paragraph()
            continue
        if l.startswith("### "):
            doc.add_heading(l[4:].strip(), level=3)
        elif l.startswith("## "):
            doc.add_heading(l[3:].strip(), level=2)
        elif l.startswith("# "):
            doc.add_heading(l[2:].strip(), level=1)
        elif l.startswith(("- ", "* ")):
            doc.add_paragraph(l[2:].strip(), style="List Bullet")
        elif l.lstrip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            doc.add_paragraph(l.lstrip(), style="List Number")
        else:
            p = doc.add_paragraph()
            # Procesar negritas **texto**
            partes = l.split("**")
            for i, t in enumerate(partes):
                run = p.add_run(t)
                if i % 2 == 1:
                    run.bold = True


def exportar_docx(
    proyecto: dict,
    especificaciones: list[dict],
    output_path: Path,
) -> Path:
    doc = Document()
    _aplicar_estilos(doc)

    # Configurar márgenes
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    # ===== Portada =====
    titulo = doc.add_heading("Especificaciones Técnicas Eléctricas", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run(proyecto.get("nombre", "—"))
    run.font.size = Pt(14)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Mandante: {proyecto.get('mandante','—')}\n").bold = True
    p.add_run(f"Ubicación: {proyecto.get('ubicacion','—')}\n")
    p.add_run(f"Tipo: {proyecto.get('tipo','—')}\n")
    p.add_run(f"Fecha: {datetime.utcnow().strftime('%d-%m-%Y')}\n")

    doc.add_page_break()

    # ===== Índice manual (Word puede generar uno real con campos TOC) =====
    doc.add_heading("Contenido", level=2)
    for i, e in enumerate(especificaciones, start=1):
        doc.add_paragraph(f"{i}. {e.get('partida','')}")
    doc.add_page_break()

    # ===== Una sección por especificación =====
    for i, e in enumerate(especificaciones, start=1):
        partida = e.get("partida", f"Partida {i}")
        doc.add_heading(f"{i}. {partida}", level=2)
        norma = e.get("norma_ref")
        if norma:
            p = doc.add_paragraph()
            p.add_run("Norma de referencia: ").bold = True
            p.add_run(str(norma))
        doc.add_paragraph()
        _agregar_contenido_md(doc, e.get("contenido_md", "") or "")
        doc.add_page_break()

    # Pie
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pie.add_run(f"Documento generado automáticamente · {datetime.utcnow().strftime('%d-%m-%Y %H:%M UTC')}")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(str(output_path))
    return output_path
