"""
Parser de Word (DOCX) usando python-docx.
Extrae párrafos, tablas y secciones con headings para identificar
estructura de bases técnicas y especificaciones.
"""
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:
    Document = None


# Palabras clave que indican secciones típicas de bases técnicas
SECCIONES_CLAVE = [
    "alcance", "objetivo", "obligaciones", "normas aplicables", "normativa",
    "requerimientos", "materiales", "marcas", "condiciones de montaje",
    "seguridad", "protocolos", "pruebas", "documentación", "documentacion",
    "calidad", "inspección", "inspeccion", "aceptación", "aceptacion",
    "exclusiones", "restricciones", "interferencias", "entregables",
]


def parse_docx(path: Path) -> dict[str, Any]:
    if Document is None:
        return {
            "tipo": "DOCX",
            "texto_plano": "",
            "estructura": {},
            "metadatos": {},
            "advertencias": ["python-docx no instalado"],
        }

    try:
        doc = Document(str(path))
    except Exception as e:
        return {
            "tipo": "DOCX",
            "texto_plano": "",
            "estructura": {},
            "metadatos": {},
            "advertencias": [f"Error leyendo DOCX: {e}"],
        }

    parrafos: list[dict] = []
    headings: list[dict] = []
    secciones_detectadas: list[str] = []
    texto_total: list[str] = []

    # Recorrer párrafos
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue

        style = p.style.name if p.style else "Normal"
        es_heading = style.lower().startswith("heading") or style.lower().startswith("título")

        parrafos.append({
            "texto": texto,
            "style": style,
            "es_heading": es_heading,
        })
        texto_total.append(texto)

        if es_heading:
            headings.append({"nivel": style, "texto": texto})

        # Detectar secciones típicas
        texto_lower = texto.lower()
        for clave in SECCIONES_CLAVE:
            if clave in texto_lower and clave not in secciones_detectadas:
                secciones_detectadas.append(clave)

    # Tablas
    tablas: list[dict] = []
    for t_idx, tabla in enumerate(doc.tables):
        filas = []
        for row in tabla.rows:
            celdas = [c.text.strip() for c in row.cells]
            filas.append(celdas)
        tablas.append({
            "tabla_idx": t_idx,
            "filas": len(filas),
            "cols": len(filas[0]) if filas else 0,
            "datos": filas,
        })
        # Agregar contenido de tablas al texto plano para RAG
        for fila in filas:
            texto_total.append(" | ".join(fila))

    # Metadatos del core_properties
    metadatos: dict[str, Any] = {
        "headings_count": len(headings),
        "parrafos_count": len(parrafos),
        "tablas_count": len(tablas),
        "secciones_clave_detectadas": secciones_detectadas,
    }
    try:
        cp = doc.core_properties
        for attr in ("author", "title", "subject", "created", "modified", "revision"):
            v = getattr(cp, attr, None)
            if v:
                metadatos[attr] = str(v)
    except Exception:
        pass

    return {
        "tipo": "DOCX",
        "texto_plano": "\n".join(texto_total),
        "estructura": {
            "parrafos": parrafos[:1000],  # cap
            "headings": headings,
            "tablas": tablas,
        },
        "metadatos": metadatos,
        "advertencias": [],
    }
